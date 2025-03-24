import os
import pandas as pd
import re
import numpy as np
from typing import List, Dict, Any, Tuple
import logging
from pathlib import Path

# Document processing libraries
import docx
import pdfplumber
import pytesseract
import cv2
from PIL import Image

# For local LLM integration
import torch
from transformers import AutoTokenizer, AutoModelForCausalLM, pipeline

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Handles extraction of text from different document formats"""
    
    @staticmethod
    def extract_from_docx(file_path: str) -> str:
        """Extract text from a DOCX file"""
        try:
            doc = docx.Document(file_path)
            full_text = []
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                full_text.append(para.text)
                
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text.append(cell.text)
                        
            return '\n'.join(full_text)
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {str(e)}")
            return ""

    @staticmethod
    def extract_from_pdf(file_path: str) -> str:
        """Extract text from a PDF file"""
        try:
            text = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    extracted = page.extract_text()
                    if extracted:
                        text.append(extracted)
            
            # If text extraction failed (possibly a scanned PDF), apply OCR
            if not ''.join(text).strip():
                logger.info(f"No text found in PDF {file_path}, attempting OCR")
                return DocumentProcessor.ocr_pdf(file_path)
                
            return '\n'.join(text)
        except Exception as e:
            logger.error(f"Error extracting text from PDF {file_path}: {str(e)}")
            return ""

    @staticmethod
    def ocr_pdf(file_path: str) -> str:
        """Apply OCR to a PDF file"""
        try:
            # Convert PDF to images and apply OCR
            from pdf2image import convert_from_path
            
            pages = convert_from_path(file_path, 300)
            text_results = []
            
            for i, page in enumerate(pages):
                text = pytesseract.image_to_string(page)
                text_results.append(text)
                
            return '\n'.join(text_results)
        except Exception as e:
            logger.error(f"Error performing OCR on PDF {file_path}: {str(e)}")
            return ""

    @staticmethod
    def extract_from_image(file_path: str) -> str:
        """Extract text from an image file using OCR"""
        try:
            # Load image and preprocess for better OCR results
            image = cv2.imread(file_path)
            if image is None:
                raise ValueError(f"Could not load image from {file_path}")
                
            # Convert to grayscale
            gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
            
            # Apply threshold to get black and white image
            _, thresh = cv2.threshold(gray, 150, 255, cv2.THRESH_BINARY | cv2.THRESH_OTSU)
            
            # Convert the OpenCV image to PIL format for pytesseract
            pil_img = Image.fromarray(thresh)
            
            # Apply OCR
            text = pytesseract.image_to_string(pil_img)
            return text
        except Exception as e:
            logger.error(f"Error extracting text from image {file_path}: {str(e)}")
            return ""

    @staticmethod
    def process_document(file_path: str) -> str:
        """Process document based on file extension"""
        file_path = file_path.strip()
        ext = os.path.splitext(file_path)[1].lower()
        
        if not os.path.exists(file_path):
            logger.error(f"File does not exist: {file_path}")
            return ""
        
        if ext == '.pdf':
            return DocumentProcessor.extract_from_pdf(file_path)
        elif ext == '.docx':
            return DocumentProcessor.extract_from_docx(file_path)
        elif ext in ['.jpg', '.jpeg', '.png', '.tiff', '.bmp']:
            return DocumentProcessor.extract_from_image(file_path)
        else:
            logger.error(f"Unsupported file format: {ext}")
            return ""


class LocalLLMProcessor:
    """Handles processing using a local LLM"""
    
    def __init__(self, model_name: str = "mistralai/Mistral-7B-Instruct-v0.2", use_cpu: bool = True):
        """Initialize the LLM processor with a specified model"""
        try:
            logger.info(f"Loading model: {model_name}")
            self.tokenizer = AutoTokenizer.from_pretrained(model_name)
            
            # Load model with settings appropriate for CPU or GPU
            if use_cpu:
                logger.info("Loading model in CPU mode")
                # Simple CPU configuration without quantization
                self.model = AutoModelForCausalLM.from_pretrained(
                    model_name,
                    device_map="auto",
                    torch_dtype=torch.float32  # Use float32 for CPU
                )
            else:
                # Original GPU configuration with quantization
                logger.info("Loading model in GPU mode with quantization")
                self.model = AutoModelForCausalLM.from_pretrained(
                    model_name,
                    torch_dtype=torch.float16,
                    device_map="auto",
                    load_in_8bit=True  # Quantize model for memory efficiency
                )
            
            self.pipe = pipeline(
                "text-generation",
                model=self.model,
                tokenizer=self.tokenizer,
                max_length=2048,
                do_sample=False
            )
            
            logger.info("Model loaded successfully")
        except Exception as e:
            logger.error(f"Error loading LLM model: {str(e)}")
            raise

    def get_completion(self, prompt: str) -> str:
        """Get a completion from the local LLM"""
        try:
            # Format prompt for instruction-tuned models
            formatted_prompt = f"<s>[INST] {prompt} [/INST]"
            
            # Generate response
            response = self.pipe(formatted_prompt)[0]['generated_text']
            
            # Extract the actual response (after the instruction)
            response = response.split('[/INST]')[-1].strip()
            
            return response
        except Exception as e:
            logger.error(f"Error generating LLM completion: {str(e)}")
            return ""


class ElementExtractor:
    """Handles the extraction of elements based on instructions"""
    
    def __init__(self, llm_processor: LocalLLMProcessor):
        """Initialize with a LocalLLMProcessor"""
        self.llm = llm_processor
        self.element_handlers = {
            'account_id': self.extract_account_id,
            'residency': self.extract_residency,
            # Add more specific handlers for different element types
        }

    def extract_element(self, element_name: str, instructions: str, document_text: str) -> Dict[str, Any]:
        """Extract element from document text based on instructions"""
        # Clean element name (lowercase, remove spaces)
        clean_element = element_name.lower().replace(' ', '_')
        
        # If there's a specific handler for this element type, use it
        if clean_element in self.element_handlers:
            return self.element_handlers[clean_element](instructions, document_text)
        
        # Otherwise use the generic LLM approach
        return self.generic_extract(element_name, instructions, document_text)
    
    def generic_extract(self, element_name: str, instructions: str, document_text: str) -> Dict[str, Any]:
        """Generic extraction using LLM"""
        # Prepare prompt for the LLM
        prompt = f"""
        I need to extract the '{element_name}' from the following document based on these instructions:
        
        INSTRUCTIONS: {instructions}
        
        DOCUMENT TEXT:
        {document_text[:1000]}...  # Truncated for LLM context limit
        
        Please extract the '{element_name}' value. If multiple values are found, list them all.
        If you cannot find the value, explain why.
        Format your response as:
        FOUND: Yes/No
        VALUE: [extracted value or N/A]
        EXPLANATION: [brief explanation of how you found it or why you couldn't find it]
        """
        
        # Get response from LLM
        response = self.llm.get_completion(prompt)
        
        # Parse LLM response
        found = "Yes" in response.split("FOUND:")[1].split("\n")[0] if "FOUND:" in response else False
        
        value = "N/A"
        if "VALUE:" in response:
            value_text = response.split("VALUE:")[1].split("\n")[0].strip()
            if value_text != "N/A":
                value = value_text
                
        explanation = response.split("EXPLANATION:")[1].strip() if "EXPLANATION:" in response else "No explanation provided"
        
        return {
            "found": found,
            "value": value,
            "explanation": explanation
        }
        
    def extract_account_id(self, instructions: str, document_text: str) -> Dict[str, Any]:
        """Extract account ID based on instructions"""
        # First try regex patterns for common account ID formats
        common_patterns = [
            r'(?:account|acct)(?:\s+|\:|\.|\#)?\s*(?:number|num|no|id|code)?(?:\s+|\:|\.|\#)?\s*(\d[\d\-]{5,})',
            r'(?:a/c|a/n)(?:\s+|\:|\.|\#)?\s*(\d[\d\-]{5,})',
            r'(?<=\s|^)(\d{5,})(?=\s|$)'  # Any 5+ digit number that stands alone
        ]
        
        # Check each pattern
        for pattern in common_patterns:
            matches = re.findall(pattern, document_text, re.IGNORECASE)
            if matches:
                # Clean up the first match (remove spaces, dashes, etc.)
                clean_id = re.sub(r'[^\d]', '', matches[0])
                if len(clean_id) >= 5:  # Assuming account IDs are at least 5 digits
                    return {
                        "found": True,
                        "value": clean_id,
                        "explanation": "Found using pattern matching for account ID format"
                    }
        
        # If regex fails, fall back to LLM
        return self.generic_extract("account_id", instructions, document_text)
        
    def extract_residency(self, instructions: str, document_text: str) -> Dict[str, Any]:
        """Extract residency based on entity type and instructions"""
        # First determine the entity type
        entity_type_prompt = f"""
        Based on the following document, determine the entity type (e.g., bank, fund, corporation, LLC, individual):
        
        DOCUMENT TEXT:
        {document_text[:1000]}...
        
        Please respond with only the entity type.
        """
        
        entity_type = self.llm.get_completion(entity_type_prompt).strip().lower()
        
        # Now extract residency based on entity type and instructions
        residency_prompt = f"""
        I need to determine the residency of a {entity_type} from this document based on these instructions:
        
        INSTRUCTIONS: {instructions}
        
        DOCUMENT TEXT:
        {document_text[:1000]}...
        
        For a {entity_type}, residency is typically determined by:
        - If it's a bank: location of headquarters or incorporation
        - If it's a fund: jurisdiction of registration or tax domicile
        - If it's an individual: tax residency, permanent address, or citizenship
        - If it's a corporation: country of incorporation or principal place of business
        
        Format your response as:
        ENTITY_TYPE: [type]
        RESIDENCY: [country/jurisdiction]
        CONFIDENCE: [High/Medium/Low]
        EXPLANATION: [brief explanation]
        """
        
        response = self.llm.get_completion(residency_prompt)
        
        # Parse LLM response to extract structured data
        residency = "Unknown"
        confidence = "Low"
        explanation = "Could not determine with confidence"
        
        if "RESIDENCY:" in response:
            residency_text = response.split("RESIDENCY:")[1].split("\n")[0].strip()
            if residency_text not in ["Unknown", "N/A", ""]:
                residency = residency_text
                
        if "CONFIDENCE:" in response:
            confidence = response.split("CONFIDENCE:")[1].split("\n")[0].strip()
            
        if "EXPLANATION:" in response:
            explanation = response.split("EXPLANATION:")[1].strip()
            
        return {
            "found": residency != "Unknown",
            "value": residency,
            "confidence": confidence,
            "entity_type": entity_type,
            "explanation": explanation
        }


class DataProcessor:
    """Main class to process the Excel data and coordinate extraction"""
    
    def __init__(self, llm_processor: LocalLLMProcessor):
        self.document_processor = DocumentProcessor()
        self.extractor = ElementExtractor(llm_processor)
        
    def process_excel(self, excel_path: str) -> Dict[str, Any]:
        """Process the Excel file and extract elements from documents"""
        try:
            # Read Excel file
            df = pd.read_excel(excel_path)
            
            # Get the directory of the Excel file to use as the base path
            excel_dir = os.path.dirname(os.path.abspath(excel_path))
            
            # Check if required columns exist
            required_columns = ['Element name', 'Instructions', 'Documents list']
            missing_columns = [col for col in required_columns if col not in df.columns]
            
            if missing_columns:
                raise ValueError(f"Excel file missing required columns: {missing_columns}")
            
            # Process each row
            results = []
            for idx, row in df.iterrows():
                element_name = row['Element name']
                instructions = row['Instructions']
                
                # Split document list (assuming comma-separated)
                doc_list_raw = [doc.strip() for doc in str(row['Documents list']).split(',')]
                
                # Convert relative paths to absolute paths
                doc_list = self.resolve_document_paths(doc_list_raw, excel_dir)
                
                element_result = self.process_element(element_name, instructions, doc_list)
                results.append(element_result)
                
            return {
                "success": True,
                "results": results
            }
            
        except Exception as e:
            logger.error(f"Error processing Excel file: {str(e)}")
            return {
                "success": False,
                "error": str(e)
            }
    
    def resolve_document_paths(self, doc_list: List[str], base_dir: str) -> List[str]:
        """Resolve document paths by searching in the base directory and its subdirectories"""
        resolved_paths = []
        
        for doc_name in doc_list:
            # Skip empty entries
            if not doc_name.strip():
                continue
                
            # If already an absolute path, use it directly
            if os.path.isabs(doc_name) and os.path.exists(doc_name):
                resolved_paths.append(doc_name)
                continue
                
            # Try direct path relative to Excel file
            direct_path = os.path.join(base_dir, doc_name)
            if os.path.exists(direct_path):
                resolved_paths.append(direct_path)
                continue
                
            # Search in subdirectories
            found = False
            for root, dirs, files in os.walk(base_dir):
                if doc_name in files:
                    full_path = os.path.join(root, doc_name)
                    resolved_paths.append(full_path)
                    found = True
                    logger.info(f"Found document '{doc_name}' at: {full_path}")
                    break
                    
            if not found:
                logger.warning(f"Could not find document: {doc_name}")
                # Add the unresolved path - the document processor will handle the error
                resolved_paths.append(doc_name)
                
        return resolved_paths
    
    def process_element(self, element_name: str, instructions: str, document_list: List[str]) -> Dict[str, Any]:
        """Process a single element across multiple documents"""
        document_results = []
        
        for doc_path in document_list:
            # Extract text from document
            logger.info(f"Processing document: {doc_path} for element: {element_name}")
            document_text = self.document_processor.process_document(doc_path)
            
            if not document_text:
                document_results.append({
                    "document_path": doc_path,
                    "success": False,
                    "error": "Failed to extract text from document"
                })
                continue
                
            # Extract element from document text
            extraction_result = self.extractor.extract_element(element_name, instructions, document_text)
            
            document_results.append({
                "document_path": doc_path,
                "success": True,
                "extraction": extraction_result
            })
            
        return {
            "element_name": element_name,
            "instructions": instructions,
            "document_results": document_results
        }


class ResultsExporter:
    """Handles exporting results to Excel or other formats"""
    
    @staticmethod
    def export_to_excel(results: Dict[str, Any], output_path: str) -> bool:
        """Export results to Excel file"""
        try:
            # Create a flattened dataframe from the results
            rows = []
            
            for element_result in results.get("results", []):
                element_name = element_result.get("element_name", "")
                instructions = element_result.get("instructions", "")
                
                for doc_result in element_result.get("document_results", []):
                    doc_path = doc_result.get("document_path", "")
                    success = doc_result.get("success", False)
                    
                    if success and "extraction" in doc_result:
                        extraction = doc_result["extraction"]
                        found = extraction.get("found", False)
                        value = extraction.get("value", "N/A")
                        explanation = extraction.get("explanation", "")
                        confidence = extraction.get("confidence", "N/A")
                        
                        rows.append({
                            "Element Name": element_name,
                            "Instructions": instructions,
                            "Document Path": doc_path,
                            "Extraction Success": "Yes" if found else "No",
                            "Extracted Value": value,
                            "Confidence": confidence if "confidence" in extraction else "N/A",
                            "Explanation": explanation
                        })
                    else:
                        error = doc_result.get("error", "Unknown error")
                        rows.append({
                            "Element Name": element_name,
                            "Instructions": instructions,
                            "Document Path": doc_path,
                            "Extraction Success": "No",
                            "Extracted Value": "N/A",
                            "Confidence": "N/A",
                            "Explanation": error
                        })
            
            # Create dataframe and export
            results_df = pd.DataFrame(rows)
            results_df.to_excel(output_path, index=False)
            
            logger.info(f"Results exported to: {output_path}")
            return True
            
        except Exception as e:
            logger.error(f"Error exporting results to Excel: {str(e)}")
            return False


def main():
    """Main function to run the document extraction system"""
    try:
        # Set up command line argument parser
        import argparse
        parser = argparse.ArgumentParser(description='Intelligent Document Element Extractor')
        parser.add_argument('--input', '-i', required=True, help='Path to input Excel file')
        parser.add_argument('--output', '-o', required=True, help='Path to output Excel file')
        parser.add_argument('--model', '-m', default="mistralai/Mistral-7B-Instruct-v0.2", 
                            help='Name of the HuggingFace model to use')
        parser.add_argument('--cpu', action='store_true', help='Force CPU mode (no CUDA/quantization)')
        
        args = parser.parse_args()
        
        # Check if CUDA is available
        cuda_available = torch.cuda.is_available()
        use_cpu = args.cpu or not cuda_available
        
        if use_cpu:
            logger.info("Running in CPU mode (no CUDA/quantization)")
        else:
            logger.info("CUDA is available, running with GPU acceleration")
        
        # Initialize LLM processor
        logger.info(f"Initializing LLM processor with model: {args.model}")
        llm_processor = LocalLLMProcessor(model_name=args.model, use_cpu=use_cpu)
        
        # Initialize data processor
        logger.info("Initializing data processor")
        processor = DataProcessor(llm_processor)
        
        # Process Excel file
        logger.info(f"Processing Excel file: {args.input}")
        results = processor.process_excel(args.input)
        
        if not results["success"]:
            logger.error(f"Failed to process Excel file: {results.get('error', 'Unknown error')}")
            return 1
            
        # Export results
        logger.info(f"Exporting results to: {args.output}")
        export_success = ResultsExporter.export_to_excel(results, args.output)
        
        if not export_success:
            logger.error("Failed to export results")
            return 1
            
        logger.info("Processing completed successfully")
        return 0
        
    except Exception as e:
        logger.error(f"Unhandled exception in main: {str(e)}")
        return 1


if __name__ == "__main__":
    exit_code = main()
    exit(exit_code)
